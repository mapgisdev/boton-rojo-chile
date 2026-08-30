"""
src/docs_generator/generate_all_chapters_docx.py
Generador automático de los 9 Capítulos Oficiales de BR-HR en formato Microsoft Word (.docx).
Diseñado con formato institucional, tipografías profesionales, tablas estilizadas, fórmulas y citas.
"""

import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

OUT_DIR = Path(__file__).resolve().parents[2] / "docs" / "capitulos_word"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# Paleta Institucional BR-HR
COLOR_PRIMARY = RGBColor(155, 34, 38)     # #9b2226 (Rojo Carmesí Oficial)
COLOR_SECONDARY = RGBColor(29, 53, 87)    # #1d3557 (Azul Marino Oscuro)
COLOR_ACCENT = RGBColor(180, 83, 9)       # #b45309 (Ámbar / Alerta Amarilla)
COLOR_DARK = RGBColor(33, 37, 41)         # #212529 (Texto Principal)
COLOR_MUTED = RGBColor(108, 117, 125)     # #6c757d (Texto Secundario)
HEX_BG_LIGHT = "F8F9FA"
HEX_BG_ACCENT = "F1FAEE"
HEX_BORDER = "CCCCCC"

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_base_doc(title_text, subtitle_text, chapter_num):
    doc = Document()
    
    # Configuración de página: Carta / Márgenes estándar 2.5 cm
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header y Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"BR-HR 2026 — Metodología Botón Rojo de Alta Resolución | Capítulo {chapter_num}")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Documento de Trabajo Técnico y Científico — Uso Oficial y Académico")
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = COLOR_MUTED

    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_DARK

    # Título del Capítulo
    p_chap = doc.add_paragraph()
    p_chap.paragraph_format.space_before = Pt(10)
    p_chap.paragraph_format.space_after = Pt(2)
    r_chap = p_chap.add_run(f"CAPÍTULO {chapter_num}")
    r_chap.bold = True
    r_chap.font.size = Pt(12)
    r_chap.font.color.rgb = COLOR_ACCENT

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title_text)
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run(subtitle_text)
    r_sub.italic = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = COLOR_SECONDARY

    # Línea divisoria
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = RGBColor(220, 220, 220)

    return doc

def add_heading1(doc, text):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = COLOR_SECONDARY
    return h

def add_heading2(doc, text):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = COLOR_PRIMARY
    return h

def add_callout(doc, text, title="NOTA TÉCNICA / PRINCIPIO METODOLÓGICO"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, HEX_BG_ACCENT)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(f"📌 {title}\n")
    r_t.bold = True
    r_t.font.size = Pt(10.5)
    r_t.font.color.rgb = COLOR_SECONDARY
    
    r_body = p.add_run(text)
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = COLOR_DARK
    
    p_post = doc.add_paragraph()
    p_post.paragraph_format.space_before = Pt(4)
    p_post.paragraph_format.space_after = Pt(4)

print("Módulo de estilos base de Word inicializado.")
